from django.shortcuts import redirect, render
from django.shortcuts import render
from django.urls import reverse
from syllabuses.models import *
from .forms import *
from django.views import View
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth import authenticate, login, logout
from django.http import HttpResponseRedirect, JsonResponse
from django.shortcuts import get_object_or_404
from django.contrib import messages
from django.contrib.admin.views.decorators import staff_member_required
from django.http import FileResponse
from docx import Document
from reportlab.pdfgen import canvas
import io


def download_syllabus_as_word(request, syllabus_id):
    syllabus = get_object_or_404(Syllabus, pk=syllabus_id)

    document = Document()
    document.add_heading(syllabus.syllabus_name, level=1)
    document.add_paragraph(f"Дисциплина: {syllabus.course}")
    document.add_paragraph(f"Уровень обучения: {syllabus.training_level}")
    document.add_paragraph(f"Язык обучения: {syllabus.language_of_education}")
    document.add_paragraph(f"Уровень владения языком: {syllabus.proficiency_level}")
    document.add_paragraph(f"Всего часов: {syllabus.total_hours}")
    document.add_paragraph(f"Классных часов: {syllabus.classroom_hours}")
    document.add_paragraph(f"Семестр: {syllabus.semester}")
    document.add_paragraph(f"ECTS кредиты: {syllabus.ects}")
    document.add_paragraph(f"СРОП часов: {syllabus.iw_hours}")
    document.add_paragraph(f"Пререквизиты: {syllabus.prerequisites}")
    document.add_paragraph(f"Формат обучения: {syllabus.format_of_training}")
    document.add_paragraph(f"Образовательные программы: {syllabus.edu_programms}")
    document.add_paragraph(f"Время и место проведения: {syllabus.time_place}")
    document.add_paragraph(f"Инструктор/Преподаватель: {syllabus.instructor}")
    document.add_paragraph(f"Цель курса: {syllabus.course_objective}")
    document.add_paragraph(f"Философия курса: {syllabus.course_philosophy}")
    document.add_paragraph(f"Политика курса: {syllabus.course_etics}")

    document.add_heading("Литература", level=2)
    for literature in syllabus.literature_set.all():
        document.add_paragraph(literature.title)

    document.add_heading("Модули", level=2)
    for module in syllabus.module_set.all():
        document.add_heading(f"Неделя {module.week}", level=3)
        document.add_paragraph(f"Тема: {module.theme}")
        document.add_paragraph(f"Формат: {module.format}")
        document.add_paragraph(f"Задания: {module.tasks}")
        document.add_paragraph(f"Результаты обучения: {module.course_lo}")
        document.add_paragraph(f"Вопросы по модулю: {module.questions}")
        document.add_paragraph(f"Оценивание: {module.grading}")
        document.add_paragraph(f"Максимальный процент: {module.max_percent}")
        document.add_paragraph(f"Максимальный вес: {module.max_weight}")
        document.add_paragraph(f"В баллах: {module.total_in_points}")

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    filename = f"{syllabus.syllabus_name}.docx"
    response = FileResponse(file_stream, as_attachment=True, filename=filename)
    return response

# def download_syllabus_as_pdf(request, syllabus_id):
#     syllabus = get_object_or_404(Syllabus, pk=syllabus_id)

#     buffer = io.BytesIO()
#     p = canvas.Canvas(buffer)

#     p.setFont("Helvetica", 12)
#     p.drawString(100, 700, syllabus.syllabus_name)
#     p.drawString(100, 650, f"Discipline: {syllabus.course}")
#     p.drawString(100, 600, f"Language of instruction: {syllabus.language_of_education}")
#     p.drawString(100, 550, f"Purpose of the course: {syllabus.course_objective}")

#     y = 500
#     literature_set = syllabus.literature_set.all()
#     for literature in literature_set:
#         p.drawString(100, y, literature.title)
#         y -= 20

#     p.showPage()
#     p.save()

#     buffer.seek(0)

#     response = FileResponse(buffer, as_attachment=True, filename=f"{syllabus.syllabus_name}.pdf")
#     return response


def download_syllabus_as_pdf(request, syllabus_id):
    syllabus = get_object_or_404(Syllabus, pk=syllabus_id)

    buffer = io.BytesIO()
    p = canvas.Canvas(buffer)

    p.setFont("Helvetica", 12)
    y = 650

    # Helper function to draw text with specified position and move the y-coordinate
    def draw_text(x, text):
        nonlocal y
        p.drawString(x, y, text)
        y -= 20

    draw_text(100, syllabus.syllabus_name)
    draw_text(100, f"Discipline: {syllabus.course}")
    draw_text(100, f"Training Level: {syllabus.training_level}")
    draw_text(100, f"Language of Education: {syllabus.language_of_education}")
    draw_text(100, f"Language Proficiency Level: {syllabus.proficiency_level}")
    draw_text(100, f"Total Hours: {syllabus.total_hours}")
    draw_text(100, f"Classroom Hours: {syllabus.classroom_hours}")
    draw_text(100, f"Semester: {syllabus.semester}")
    draw_text(100, f"ECTS Credits: {syllabus.ects}")
    draw_text(100, f"IW Hours: {syllabus.iw_hours}")
    draw_text(100, f"Prerequisites: {syllabus.prerequisites}")
    draw_text(100, f"Training Format: {syllabus.format_of_training}")
    draw_text(100, f"Educational Programs: {syllabus.edu_programms}")
    draw_text(100, f"Time and Place of Conduct: {syllabus.time_place}")
    draw_text(100, f"Instructor/Teacher: {syllabus.instructor}")
    draw_text(100, f"Course Objective: {syllabus.course_objective}")
    draw_text(100, f"Course Philosophy: {syllabus.course_philosophy}")
    draw_text(100, f"Course Policy: {syllabus.course_etics}")

    draw_text(100, "Literature")
    for literature in syllabus.literature_set.all():
        draw_text(120, literature.title)

    draw_text(100, "Modules")
    for module in syllabus.module_set.all():
        draw_text(120, f"Week {module.week}")
        draw_text(120, f"Theme: {module.theme}")
        draw_text(120, f"Format: {module.format}")
        draw_text(120, f"Tasks: {module.tasks}")
        draw_text(120, f"Course Learning Outcomes: {module.course_lo}")
        draw_text(120, f"Module Questions: {module.questions}")
        draw_text(120, f"Grading: {module.grading}")
        draw_text(120, f"Maximum Percentage: {module.max_percent}")
        draw_text(120, f"Maximum Weight: {module.max_weight}")
        draw_text(120, f"In Points: {module.total_in_points}")

    p.showPage()
    p.save()

    buffer.seek(0)

    response = FileResponse(buffer, as_attachment=True, filename=f"{syllabus.syllabus_name}.pdf")
    return response


# def download_syllabus_as_pdf(request, syllabus_id):
#     syllabus = get_object_or_404(Syllabus, pk=syllabus_id)

#     buffer = io.BytesIO()
#     p = canvas.Canvas(buffer)

    # p.setFont("Helvetica", 12)
    # p.drawString(100, 700, syllabus.syllabus_name)
    # p.drawString(100, 650, f"Discipline: {syllabus.course}")
    # p.drawString(100, 600, f"Language of instruction: {syllabus.language_of_education}")
    # p.drawString(100, 550, f"Purpose of the course: {syllabus.course_objective}")

    
    # p.drawString(syllabus.syllabus_name)
    # p.drawString(100, 700, f"Дисциплина: {syllabus.course}")
    # p.drawString(100, 700, f"Уровень обучения: {syllabus.training_level}")
    # p.drawString(100, 700, f"Язык обучения: {syllabus.language_of_education}")
    # p.drawString(100, 700, f"Уровень владения языком: {syllabus.proficiency_level}")
    # p.drawString(100, 700, f"Всего часов: {syllabus.total_hours}")
    # p.drawString(100, 700, f"Классных часов: {syllabus.classroom_hours}")
    # p.drawString(100, 700, f"Семестр: {syllabus.semester}")
    # p.drawString(100, 700, f"ECTS кредиты: {syllabus.ects}")
    # p.drawString(100, 700, f"СРОП часов: {syllabus.iw_hours}")
    # p.drawString(100, 700, f"Пререквизиты: {syllabus.prerequisites}")
    # p.drawString(100, 700, f"Образовательные программы: {syllabus.edu_programms}")
    # p.drawString(100, 700, f"Формат обучения: {syllabus.format_of_training}")
    # p.drawString(100, 700, f"Время и место проведения: {syllabus.time_place}")
    # p.drawString(100, 700, f"Инструктор/Преподаватель: {syllabus.instructor}")
    # p.drawString(100, 700, f"Цель курса: {syllabus.course_objective}")
    # p.drawString(100, 700, f"Философия курса: {syllabus.course_philosophy}")
    # p.drawString(100, 700, f"Политика курса: {syllabus.course_etics}")

    # p.drawString("Литература", level=2)
    # for literature in syllabus.literature_set.all():
    #     p.drawString(literature.title)

    # p.drawString("Модули", level=2)
    # for module in syllabus.module_set.all():
    #     p.drawString(f"Неделя {module.week}", level=3)
    #     p.drawString(f"Тема: {module.theme}")
    #     p.drawString(f"Формат: {module.format}")
    #     p.drawString(f"Задания: {module.tasks}")
    #     p.drawString(f"Результаты обучения: {module.course_lo}")
    #     p.drawString(f"Вопросы по модулю: {module.questions}")
    #     p.drawString(f"Оценивание: {module.grading}")
    #     p.drawString(f"Максимальный процент: {module.max_percent}")
    #     p.drawString(f"Максимальный вес: {module.max_weight}")
    #     p.drawString(f"В баллах: {module.total_in_points}")

    # y = 500
    # literature_set = syllabus.literature_set.all()
    # for literature in literature_set:
    #     p.drawString(100, y, literature.title)
    #     y -= 20

    # p.showPage()
    # p.save()

    # buffer.seek(0)

    # response = FileResponse(buffer, as_attachment=True, filename=f"{syllabus.syllabus_name}.pdf")
    # return response
# Create your views here.
def home(request):
    return render(request, 'syllabuses/home.html', {})



def login_v(request):
    if request.method == 'POST':
        form = AuthenticationForm(data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(request, username=username, password=password)
            if user is not None:
                login(request, user)
                return redirect('create_syllabus')  # Исправленный путь для перенаправления
            else:
                print("Why is this not returned for inval")
    else:
        form = AuthenticationForm()
    return render(request, 'syllabuses/login.html', {'form': form})


def logout_view(request):
    logout(request)
    return redirect('../create_syllabus')






def create_syllabus(request):
    if request.method == 'POST':
        form = SyllabusForm(request.POST)
        if form.is_valid():
            syllabus = form.save(commit=False)
            syllabus.status = Status.objects.get(type="Created")
            syllabus.save()
            return redirect(f'literature_form/{syllabus.id}')
    else:
        form = SyllabusForm()
    return render(request, 'syllabuses/create_syllabus.html', {'form': form})



def next_step(request, syllabus_id: int):
    syllabus = Syllabus.objects.get(pk=syllabus_id)
    if request.method == "POST":
        syllabus.status = Status.objects.get(type="added literature")
        syllabus.save()
        l = request.POST["liter"]
        literature = Literature.objects.get(pk=l)
        mandatory = request.POST["mandatory"]
        LiteratureInSyllabus.objects.create(
            syllabus = syllabus,
            literature = literature,
            mandatory = mandatory
        )
    return render(request, 'syllabuses/next_step.html',{
                      'syllabus': syllabus, 
                      'literatures': Literature.objects.filter(course=syllabus.course),
                      'literaturesinsyllabus': LiteratureInSyllabus.objects.filter(syllabus = syllabus)
                    })

def delete_literature(request, pk, syllabus_id):
    literature = LiteratureInSyllabus.objects.get(pk=pk)
    syllabus_id=syllabus_id
    literature.delete()
    return redirect(f'../../../next_step/{syllabus_id}')


def delete_syllabus(request, syllabus_id):
    syl = Syllabus.objects.get(pk=syllabus_id)
    syl.delete()
    return redirect(f'../../../my_syllabuses')

def delete_module(request, pk, syllabus_id):
    module = Module.objects.get(pk=pk)
    syllabus_id=syllabus_id
    module.delete()
    return redirect(f'../../../add_module/{syllabus_id}')

def add_literature(request, syllabus_id: int):
    syllabus = Syllabus.objects.get(pk=syllabus_id)
    # literature_form = SecondStepForm()
    if request.method == "POST":
        title = request.POST["title"]
        Literature.objects.create(
            course=syllabus.course,
            title=title,
        )


    return render(request, 'syllabuses/literature_form.html', 
                  {
                      'syllabus': syllabus, 
                      'literatures': Literature.objects.filter(course=syllabus.course),
                    })



def syllabus_details(request, syllabus_id: int):
    syllabus = get_object_or_404(Syllabus, pk=syllabus_id)
    lo11 = CourseLO.objects.filter(syllabus=syllabus).filter(type=True)
    lo22 = CourseLO.objects.filter(syllabus=syllabus).filter(type=False)
    modules = Module.objects.filter(syllabus=syllabus).order_by('week')

    if request.method == 'POST':
        syllabus.syllabus_name= request.POST.get('syllabus_name')
        course = request.POST.get('course')
        training_level = request.POST.get('training_level')
        language_of_education = request.POST.get('language_of_education')
        proficiency_level = request.POST.get('proficiency_level')
        total_hours = request.POST.get('total_hours')
        classroom_hours = request.POST.get('classroom_hours')
        semester = request.POST.get('semester')
        ects = request.POST.get('ects')
        iw_hours = request.POST.get('iw_hours')
        prerequisites = request.POST.get('prerequisites')
        format_of_training = request.POST.get('format_of_training')
        edu_programms = request.POST.get('edu_programms')
        time_place = request.POST.get('time_place')
        instructor = request.POST.get('instructor')
        course_objective = request.POST.get('course_objective')
        agreed_with = request.POST.get('agreed_with')
        status = request.POST.get('status')
        course_philosophy = request.POST.get('course_philosophy')
        course_etics = request.POST.get('course_etics')
        asu = request.POST.get('asu')
        syllabus.course = Course.objects.get(pk=course)
        syllabus.training_level = EduLevel.objects.get(pk=training_level)
        syllabus.language_of_education = Language.objects.get(pk=language_of_education)
        syllabus.proficiency_level = Proficiency.objects.get(pk=proficiency_level)
        syllabus.total_hours = total_hours
        syllabus.classroom_hours = classroom_hours
        syllabus.semester = semester
        syllabus.ects = ects
        syllabus.iw_hours = iw_hours
        syllabus.prerequisites = prerequisites
        syllabus.format_of_training = Format.objects.get(pk=format_of_training)
        syllabus.edu_programms = edu_programms
        syllabus.time_place = time_place
        syllabus.instructor = CustomUser.objects.get(pk=instructor)
        syllabus.course_objective = course_objective
        syllabus.agreed_with = Director.objects.get(pk=agreed_with)
        syllabus.status = Status.objects.get(pk=status)
        syllabus.course_philosophy = course_philosophy
        syllabus.course_etics = course_etics
        if asu=="on":
            syllabus.asu = True
        else: 
            syllabus.asu = False
        syllabus.save()
        return redirect(reverse('syllabus_details', kwargs={'syllabus_id': syllabus_id}))

    
    
    return render(request, 'syllabuses/syllabus_details.html', {
        'syllabus': syllabus,
        'literatures': LiteratureInSyllabus.objects.filter(syllabus=syllabus),
        'lo11': lo11,
        'lo22': lo22,
        'modules': modules,
        'courses': Course.objects.all(),
        'edu_levels': EduLevel.objects.all(),
        'languages': Language.objects.all(),
        'proficiencies': Proficiency.objects.all(),
        'formats': Format.objects.all(),
        'instructors': CustomUser.objects.all(),
        'directors': Director.objects.all(),
        'statuses': Status.objects.all(),
    })





def half(request, syllabus_id: int):
    syllabus = Syllabus.objects.get(pk=syllabus_id)
  # literature_form = SecondStepForm()
    if request.method == "POST":
        syllabus.status = Status.objects.get(type="added lo")
        syllabus.save()
        lo = request.POST["lo"]
        lo2 = request.POST["lo2"]
        CourseLO.objects.create(
            syllabus=syllabus,
            type = True,
            info = lo
        )
        CourseLO.objects.create(
            syllabus=syllabus,
            type = False,
            info = lo2
        )


    return render(request, 'syllabuses/half.html', 
                  {
                      'syllabus': syllabus, 

                    })

def edit_profile(request):
    currentuser = request.user

    return render(request, 'syllabuses/edit_profile.html', 
                  {
                      'user': currentuser, 

                    })




def add_policy(request, syllabus_id: int):
    syllabus = Syllabus.objects.get(pk=syllabus_id)
    if request.method == "POST":
        syllabus.status = Status.objects.get(type="added policy")
        syllabus.course_philosophy = request.POST["phylosophy"]
        syllabus.course_etics = request.POST["policy"]
        syllabus.save()
        messages.success(request, "Силлабус создан")
        return redirect('home')  # Перенаправление на главный экран

    return render(request, 'syllabuses/add_policy.html',
                   {
                       'syllabus': syllabus
                       })



def add_module(request, syllabus_id: int):
    syllabus = Syllabus.objects.get(pk=syllabus_id)

    if request.method == "POST":
        syllabus.status = Status.objects.get(type="added modules")
        syllabus.save()
        week = request.POST["week"]
        theme = request.POST["theme"]
        format = Format.objects.get(pk = request.POST["format"]) 
        tasks = request.POST["tasks"]
        lo = request.POST["lo"]
        questions = request.POST["questions"]
        liter = LiteratureInSyllabus.objects.get(pk = request.POST["liter"]) 
        grading = request.POST["grading"]
        maxpercent = request.POST["maxpercent"]
        maxvalue = request.POST["maxvalue"]
        total_in_points = request.POST["total_in_points"]

        Module.objects.create(
            syllabus=syllabus,
            week=week,
            tasks=tasks,
            course_lo=lo,
            theme=theme,
            format=format,
            questions=questions,
            literature=liter,
            grading=grading,
            max_percent=maxpercent,
            max_weight=maxvalue,
            total_in_points=total_in_points,
        )


    return render(request, 'syllabuses/add_module.html', 
                  {
                      'syllabus': syllabus, 
                      'literatures': LiteratureInSyllabus.objects.filter(syllabus=syllabus),
                      'formats': Format.objects.all(),
                      'modules': Module.objects.filter(syllabus=syllabus).order_by('week'),
                      'lo11': CourseLO.objects.get(syllabus=syllabus, type=True),
                      'lo22': CourseLO.objects.get(syllabus=syllabus, type=False),
                    })



def continue_edit(request, syllabus_id: int):
    syllabus = Syllabus.objects.get(pk=syllabus_id)
    print(syllabus.status)

    if syllabus.status == Status.objects.get(type="Created"):
        return redirect(f'../../literature_form/{syllabus_id}')
    elif syllabus.status == Status.objects.get(type="added literature"):
        return redirect(f'../../half/{syllabus_id}')
    elif syllabus.status == Status.objects.get(type="added lo"):
        return redirect(f'../../add_module/{syllabus_id}')
    else: return redirect(reverse('syllabus_details', args=[syllabus_id]))






@staff_member_required
def my_syllabuses(request):
    syllabuses = Syllabus.objects.all()  # Получаем все силлабусы

    return render(request, 'syllabuses/my_syllabuses.html', {'syllabuses': syllabuses})



def view_syllabuses(request):
    user = request.user
    syllabuses = user.get_created_syllabuses()
    context = {
        'syllabuses': syllabuses
    }
    return render(request, 'syllabuses/view_syllabuses.html', context)

   

def add_instructor(request):
    if request.method == 'POST':
        form = CustomUserForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('../create_syllabus')  # Перенаправьте пользователя на страницу успешного добавления преподавателя
    else:
        form = CustomUserForm()
    return render(request, 'syllabuses/add_instructor.html', {'form': form})

class SchoolView(View):
    def get(self, request):
        schools = School.objects.all()
        return render(request, 'syllabuses/school_list.html', {'schools': schools})

    def post(self, request):
        form = SchoolForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('syllabuses:school_list')
        return render(request, 'syllabuses/school_form.html', {'form': form})

class CustomUserView(View):
    def get(self, request):
        users = CustomUser.objects.all()
        return render(request, 'syllabuses/user_list.html', {'users': users})

    def post(self, request):
        form = CustomUserForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('syllabuses:user_list')
        return render(request, 'syllabuses/user_form.html', {'form': form})

class DirectorView(View):
    def get(self, request):
        directors = Director.objects.all()
        return render(request, 'syllabuses/director_list.html', {'directors': directors})

    def post(self, request):
        form = DirectorForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('syllabuses:director_list')
        return render(request, 'syllabuses/director_form.html', {'form': form})